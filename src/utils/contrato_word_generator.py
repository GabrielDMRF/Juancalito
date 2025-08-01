#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generador de Contrato en Word a partir de una plantilla
"""
import os
import sys
import subprocess
from datetime import datetime, timedelta
from pathlib import Path

# Importar modelos necesarios
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from models.database import Empleado, Contrato, get_db
from views import config

def safe_filename(filename: str) -> str:
    """Crear nombre de archivo seguro (sin caracteres inválidos)"""
    invalid_chars = '<>:"/\\|?*'
    for char in invalid_chars:
        filename = filename.replace(char, '_')
    filename = ''.join(c for c in filename if c.isalnum() or c in (' ', '-', '_')).strip()
    return filename.replace(' ', '_')[:50]

def numero_a_letras(numero: int) -> str:
    """Convertir número a letras para representar montos en pesos M/CTE"""
    try:
        if numero == 1423500:
            return "UN MILLÓN CUATROCIENTOS VEINTITRÉS MIL QUINIENTOS PESOS M/CTE"
        elif numero >= 1000000:
            millones = numero // 1000000
            resto = numero % 1000000
            # Parte de millones
            if millones == 1:
                millones_txt = "UN MILLÓN"
            elif millones < 10:
                unidades = ["", "UN", "DOS", "TRES", "CUATRO", "CINCO", "SEIS", "SIETE", "OCHO", "NUEVE"]
                millones_txt = f"{unidades[millones]} MILLONES" if millones < len(unidades) else f"{millones} MILLONES"
            else:
                millones_txt = f"{millones} MILLONES"
            # Parte de miles (resto)
            if resto == 0:
                return f"{millones_txt} PESOS M/CTE"
            else:
                if resto < 1000:
                    resto_txt = str(resto)
                else:
                    miles = resto // 1000
                    rem = resto % 1000
                    if rem == 0:
                        resto_txt = f"{miles} MIL"
                    else:
                        resto_txt = f"{miles} MIL {rem}"
                return f"{millones_txt} {resto_txt} PESOS M/CTE"
        else:
            if numero < 1000:
                letras = str(numero)
            else:
                miles = numero // 1000
                rem = numero % 1000
                if rem == 0:
                    letras = f"{miles} MIL"
                else:
                    letras = f"{miles} MIL {rem}"
            return letras + " PESOS M/CTE"
    except Exception:
        return "SALARIO EN LETRAS"

def generar_contrato_word(contrato: Contrato) -> dict:
    """Generar un documento Word de contrato a partir de la plantilla y datos del contrato dado."""
    try:
        # Importar Document de python-docx
        try:
            from docx import Document
        except ImportError:
            return {
                'success': False,
                'message': "La librería 'python-docx' no está instalada."
            }
        # Obtener empleado relacionado
        empleado = contrato.empleado  # relación lazy de SQLAlchemy
        if not empleado:
            empleado = get_db().query(Empleado).filter(Empleado.id == contrato.empleado_id).first()
        if not empleado:
            return {
                'success': False,
                'message': "Empleado asociado no encontrado."
            }
        # Determinar carpeta del empleado
        nombre_seguro = safe_filename(empleado.nombre_completo)
        cedula_segura = safe_filename(empleado.cedula)
        base_dir = Path("empleados_data")
        folder_name = f"{nombre_seguro}_{cedula_segura}"
        employee_dir = base_dir / folder_name
        contratos_dir = employee_dir / "contratos"
        try:
            contratos_dir.mkdir(parents=True, exist_ok=True)
        except Exception as e:
            return {
                'success': False,
                'message': f"Error creando carpeta para contratos: {e}"
            }
        # Cargar documento plantilla
        template_path = config.RUTA_PLANTILLA_CONTRATO_WORD
        if not template_path.exists():
            return {
                'success': False,
                'message': f"Plantilla Word no encontrada en: {template_path}"
            }
        doc = Document(str(template_path))
        # Preparar datos para reemplazar
        salario_base = contrato.salario_base or 1423500  # valor por defecto si no hay salario
        subsidio = contrato.subsidio_transporte or 140606
        fecha_elaboracion = datetime.now().strftime("%d DE %B DE %Y").upper()
        fecha_inicio_str = contrato.fecha_inicio.strftime("%d DE %B %Y") if contrato.fecha_inicio else ""
        fecha_fin_str = contrato.fecha_fin.strftime("%d DE %B %Y") if contrato.fecha_fin else ""
        # Término del contrato: "A DOS MESES" para temporales, "INDEFINIDO" para permanentes
        termino_text = "A DOS MESES" if contrato.tipo_contrato == "temporal" else "INDEFINIDO"
        # Lugar de nacimiento y nacionalidad
        lugar_nac = f"{empleado.direccion}, COLOMBIANO(A)" if empleado.direccion else "XXXXXXXXXXXXXXX, COLOMBIANO(A)"
        # Salario en letras
        salario_letras = numero_a_letras(salario_base)
        # Reemplazar en párrafos fuera de tablas
        for para in doc.paragraphs:
            text = para.text
            # Nombre del trabajador (placeholder X repetidas)
            if "NOMBRE DEL TRABAJADOR" in text:
                # El nombre suele estar en el siguiente párrafo de X, se manejará más abajo
                continue
            if text.strip().startswith("XXXXXXXXXXXXXXXX"):
                # Asumimos que la primera ocurrencia de X largos es el nombre
                para.text = para.text.replace("XXXXXXXXXXXXXXXXXXXX", empleado.nombre_completo)
                para.text = para.text.replace("XXXXXXXXXXXXXXXXX", empleado.nombre_completo)
            # Dirección del trabajador
            if "DIRECCION TRABAJADOR" in text:
                continue
            if text.strip().startswith("XXXXXXXXXXXXXXX"):
                # Reemplazo genérico para dirección o lugar
                if empleado.direccion and empleado.direccion in lugar_nac:
                    # Reemplazar como dirección si coincide con lugar_nac
                    para.text = para.text.replace("XXXXXXXXXXXXXXX", empleado.direccion)
                else:
                    para.text = para.text.replace("XXXXXXXXXXXXXXX", empleado.direccion or "")
            # Cargo del trabajador
            if "OPERARIO(A)" in text or "OFICIOS VARIOS" in text:
                # Reemplazar texto de cargo por el real si existe
                if empleado.cargo:
                    para.text = empleado.cargo
            # Fecha de inicio
            if "FECHA DE INICIACION" in text or "FECHA DE INICIACIÓN" in text:
                continue
            if text.strip().endswith("DE 2025") or text.strip().endswith("DE 2026") or text.strip().endswith("DE 2024"):
                # Asumimos que un párrafo que termina con año es una fecha potencial a reemplazar
                if fecha_inicio_str and para.text.strip() == "17 DE FEBRERO 2025":
                    para.text = fecha_inicio_str
                elif fecha_fin_str and para.text.strip() == "16 DE ABRIL 2025":
                    para.text = fecha_fin_str
            # Termino inicial del contrato
            if "TERMINO INICIAL" in text or "TÉRMINO INICIAL" in text:
                continue
            if "A DOS MESES" in text or "INDEFINIDO" in text:
                para.text = termino_text
            # Salario ordinario/integral valor y valor en letras
            if para.text.strip().startswith("$"):
                # Reemplazar salario numérico (formato con puntos y punto final)
                salario_formateado = f"{salario_base:,}".replace(",", ".")
                para.text = f"$ {salario_formateado}."
            if "PESOS M/CTE" in para.text:
                # Reemplazar salario en letras (asegurando punto final)
                para.text = salario_letras
                if not para.text.endswith("."):
                    para.text += "."
        # Reemplazar en tablas (si la plantilla usa tablas para formatear campos)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        # Nombre
                        if "XXXXXXXXXXXXXXXX" in text:
                            para.text = text.replace("XXXXXXXXXXXXXXXXXXXX", empleado.nombre_completo).replace("XXXXXXXXXXXXXXXXX", empleado.nombre_completo)
                        # Dirección
                        if "XXXXXXXXXXXXXXX" in text:
                            para.text = text.replace("XXXXXXXXXXXXXXX", empleado.direccion or "")
                        # Lugar de nacimiento
                        if "XXXXXXXXXXXXXXX, COLOMBIANO" in text:
                            para.text = text.replace("XXXXXXXXXXXXXXX, COLOMBIANO(A)", lugar_nac)
                        # Cargo
                        if "OPERARIO" in text or "OFICIOS VARIOS" in text:
                            if empleado.cargo:
                                para.text = empleado.cargo
                        # Salario valor
                        if cell.text.strip().startswith("$"):
                            salario_formateado = f"{salario_base:,}".replace(",", ".")
                            para.text = f"$ {salario_formateado}."
                        # Salario en letras
                        if "PESOS M/CTE" in text and (text.strip().upper().startswith("UN ") or text.strip().endswith("PESOS M/CTE")):
                            para.text = salario_letras
                            if not para.text.endswith("."):
                                para.text += "."
                        # Fecha inicio
                        if para.text.strip() == "17 DE FEBRERO 2025" and fecha_inicio_str:
                            para.text = fecha_inicio_str
                        # Fecha fin
                        if para.text.strip() == "16 DE ABRIL 2025":
                            para.text = fecha_fin_str if fecha_fin_str else ""
                        # Termino contrato
                        if para.text.strip() == "A DOS MESES" or para.text.strip() == "INDEFINIDO":
                            para.text = termino_text
        # Generar nombre de archivo único para el contrato Word
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        contrato_filename = f"contrato_profesional_{cedula_segura}_{timestamp}.docx"
        contrato_path = contratos_dir / contrato_filename
        # Guardar documento Word
        try:
            doc.save(str(contrato_path))
        except Exception as e:
            return {
                'success': False,
                'message': f"Error al guardar el documento Word: {e}"
            }
        # Abrir automáticamente el archivo generado
        try:
            if sys.platform.startswith("win32") or sys.platform.startswith("win"):
                os.startfile(str(contrato_path))
            elif sys.platform == "darwin":
                subprocess.run(["open", str(contrato_path)], check=True)
            else:
                subprocess.run(["xdg-open", str(contrato_path)], check=True)
        except Exception as open_err:
            print(f"⚠️ Advertencia: no se pudo abrir automáticamente el archivo: {open_err}")
        return {
            'success': True,
            'file_path': str(contrato_path),
            'message': f"Contrato Word generado exitosamente: {contrato_filename}"
        }
    except Exception as e:
        return {
            'success': False,
            'message': f"Ocurrió un error inesperado: {e}"
        }
